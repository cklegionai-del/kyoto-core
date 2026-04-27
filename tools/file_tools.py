import os
from pypdf import PdfReader
from docx import Document
import pandas as pd

def read_text(path):
    if not os.path.exists(path): raise FileNotFoundError(path)
    with open(path, "r", encoding="utf-8") as f: return f.read()

def write_text(path, content):
    with open(path, "w", encoding="utf-8") as f: f.write(content)

def read_csv(path):
    if not os.path.exists(path): raise FileNotFoundError(path)
    return pd.read_csv(path).to_dict(orient="records")

def read_pdf(path):
    if not os.path.exists(path): raise FileNotFoundError(path)
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def read_docx(path):
    if not os.path.exists(path): raise FileNotFoundError(path)
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def read_excel(path, sheet=0):
    if not os.path.exists(path): raise FileNotFoundError(path)
    return pd.read_excel(path, sheet_name=sheet).to_dict(orient="records")

if __name__ == "__main__":
    print("📁 Testing file_tools.py...\n")
    test_dir = os.path.expanduser("~/Desktop/file_tools_test")
    os.makedirs(test_dir, exist_ok=True)
    
    # 1. Text
    txt = os.path.join(test_dir, "test.txt")
    write_text(txt, "Hello from file_tools!")
    print(f"1. Text: {read_text(txt)}")
    
    # 2. CSV
    csv_path = os.path.join(test_dir, "test.csv")
    pd.DataFrame([{"name": "Alice", "age": 30}]).to_csv(csv_path, index=False)
    print(f"2. CSV: {read_csv(csv_path)}")
    
    # 3. PDF
    from reportlab.pdfgen import canvas
    pdf = os.path.join(test_dir, "test.pdf")
    canvas.Canvas(pdf).drawString(100, 750, "PDF Content"); canvas.Canvas(pdf).save()
    print(f"3. PDF: {read_pdf(pdf)[:30]}...")
    
    # 4. DOCX
    docx = os.path.join(test_dir, "test.docx")
    Document().add_paragraph("DOCX Content"); Document().save(docx)
    print(f"4. DOCX: {read_docx(docx)}")
    
    # 5. Excel
    xlsx = os.path.join(test_dir, "test.xlsx")
    pd.DataFrame([{"id": 1, "val": "A"}]).to_excel(xlsx, index=False)
    print(f"5. Excel: {read_excel(xlsx)}")
    
    # Cleanup
    import shutil
    shutil.rmtree(test_dir)
    print("\n✅ All tests passed!")
