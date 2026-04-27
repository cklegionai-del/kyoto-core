import os, subprocess
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd

def read_pdf(path):
    """Extract text from PDF using pdftotext (poppler) with PyPDF2 fallback."""
    if not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}
    
    # Try pdftotext first (more reliable for complex layouts)
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", path, "-"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return {"success": True, "data": result.stdout.strip(), "source": "pdftotext"}
    except Exception as e:
        pass  # Fallback to PyPDF2
    
    # Fallback: PyPDF2
    try:
        with open(path, 'rb') as file:
            reader = PdfReader(file)
            text = ''.join(page.extract_text() or '' for page in reader.pages)
        if text.strip():
            return {"success": True, "data": text.strip(), "source": "PyPDF2"}
        return {"success": False, "error": "No text extracted (PDF may be scanned/image-based)"}
    except Exception as e:
        return {"success": False, "error": f"PyPDF2 error: {str(e)}"}

def read_docx(path):
    """Extract text from Word document."""
    if not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}
    try:
        document = Document(path)
        text = '\n'.join([para.text for para in document.paragraphs if para.text.strip()])
        return {"success": True, "data": text.strip(), "source": "python-docx"}
    except Exception as e:
        return {"success": False, "error": f"docx error: {str(e)}"}

def read_excel(path):
    """Extract data from Excel as formatted string."""
    if not os.path.exists(path):
        return {"success": False, "error": f"File not found: {path}"}
    try:
        df = pd.read_excel(path)
        return {"success": True, "data": df.to_string(), "source": "pandas"}
    except Exception as e:
        return {"success": False, "error": f"excel error: {str(e)}"}

if __name__ == "__main__":
    print("📚 Testing doc_reader.py...\n")
    
    # Find any PDF on Desktop to test
    desktop = os.path.expanduser("~/Desktop")
    test_pdf = None
    for f in os.listdir(desktop):
        if f.lower().endswith(".pdf"):
            test_pdf = os.path.join(desktop, f)
            break
    
    if test_pdf:
        print(f"📄 Testing with: {os.path.basename(test_pdf)}")
        result = read_pdf(test_pdf)
        if result["success"]:
            print(f"✅ Success! Source: {result['source']}")
            print(f"📝 Preview ({len(result['data'])} chars):\n{result['data'][:300]}...")
        else:
            print(f"❌ {result['error']}")
    else:
        print("⚠️ No PDF found on Desktop — create one or skip test")
    
    print("\n✅ Tests complete!")

# === Word (.docx) Support ===
from docx import Document
def read_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# === Excel (.xlsx) Support ===
import pandas as pd
def read_excel(path: str, sheet: str = None) -> str:
    df = pd.read_excel(path, sheet_name=sheet)
    return df.to_markdown(index=False)  # Returns clean Markdown table

# === PDF (Enhanced with pdfplumber) ===
import pdfplumber
def read_pdf_enhanced(path: str) -> str:
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt: text.append(txt)
    return "\n\n".join(text)

def read_any_file(path: str) -> str:
    """Auto-detect format and extract text"""
    from pathlib import Path
    ext = Path(path).suffix.lower()
    if ext == '.pdf': return read_pdf_enhanced(path)
    if ext == '.docx': return read_docx(path)
    if ext in ['.xlsx', '.xls']: return read_excel(path)
    if ext == '.md': return Path(path).read_text(encoding='utf-8')
    if ext == '.txt': return Path(path).read_text(encoding='utf-8')
    return f"[Unsupported format: {ext}]"
